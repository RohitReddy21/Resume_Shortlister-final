import os
import sys

# Ensure the backend project root is on sys.path so `app` package imports work
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from PIL import Image
import zipfile

from app.services.anonymizer import mask_images_in_pdf, mask_images_in_docx


def create_sample_image(path):
    im = Image.new("RGB", (200, 100), color=(255, 0, 0))
    im.save(path, format="PNG")


def test_mask_images_pdf(tmp_path):
    img_path = tmp_path / "sample.png"
    create_sample_image(str(img_path))

    pdf_path = tmp_path / "sample.pdf"
    # Save a simple PDF containing the image
    im = Image.open(str(img_path))
    im.save(str(pdf_path), "PDF", resolution=100.0)

    masked = tmp_path / "sample.masked.pdf"
    # Blur images
    mask_images_in_pdf(str(pdf_path), str(masked), remove=False, blur_radius=5)
    assert masked.exists()
    assert masked.stat().st_size > 0

    # Remove images
    masked2 = tmp_path / "sample.masked2.pdf"
    mask_images_in_pdf(str(pdf_path), str(masked2), remove=True, blur_radius=None)
    assert masked2.exists()
    assert masked2.stat().st_size > 0


def test_mask_images_docx(tmp_path):
    # create image
    img_path = tmp_path / "sample2.png"
    create_sample_image(str(img_path))

    # create a docx with the image
    from docx import Document

    doc = Document()
    doc.add_paragraph("Test doc with image")
    doc.add_picture(str(img_path))
    docx_path = tmp_path / "sample.docx"
    doc.save(str(docx_path))

    masked = tmp_path / "sample.masked.docx"
    mask_images_in_docx(str(docx_path), str(masked), remove=False, blur_radius=5)
    assert masked.exists()
    assert masked.stat().st_size > 0

    # compare media files inside original and masked
    def list_media_sizes(path):
        sizes = {}
        with zipfile.ZipFile(path, "r") as z:
            for name in z.namelist():
                if name.startswith("word/media/"):
                    sizes[name] = len(z.read(name))
        return sizes

    orig_media = list_media_sizes(str(docx_path))
    masked_media = list_media_sizes(str(masked))
    # both should have media entries
    assert orig_media
    assert masked_media
    # masked media should exist; content may differ depending on image encoding
    assert set(orig_media.keys()) == set(masked_media.keys())
